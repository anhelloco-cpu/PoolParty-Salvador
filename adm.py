import streamlit as st
import pandas as pd
from supabase import create_client, Client
from io import BytesIO
from datetime import datetime
from fpdf import FPDF
from docx import Document

# --- CONFIGURACIÓN SUPABASE ---
SUPABASE_URL = "https://auezltquejptsupqkcqh.supabase.co"
SUPABASE_KEY = "sb_publishable_eImPwr3l_Wq-TO3FW4wk2g_YUCE898x"
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# Configuración de la página
st.set_page_config(page_title="Panel VIP - Pool Party", layout="wide")

st.title("📋 Panel de Control Financiero - Pool Party")
st.markdown("Gestión de invitados y presupuesto profesional (Costos Fijos vs Variables).")

# --- FUNCIÓN PARA ELIMINAR DUPLICADOS ---
def eliminar_duplicados(df):
    df['nombre_completo'] = (df['nombre'].str.strip() + " " + df['apellido'].str.strip()).str.lower()
    duplicados_a_borrar = df[df.duplicated(subset=['nombre_completo'], keep='first')]
    ids_a_eliminar = duplicados_a_borrar['id'].tolist()
    
    exitos, errores = 0, 0
    for record_id in ids_a_eliminar:
        try:
            supabase.table("invitados_pool_party").delete().eq("id", record_id).execute()
            exitos += 1
        except:
            errores += 1
            
    return exitos, errores

# --- INICIALIZAR MEMORIA DE APORTANTES ---
if 'aportantes' not in st.session_state:
    st.session_state.aportantes = []

# --- CARGA DE DATOS ---
try:
    respuesta = supabase.table("invitados_pool_party").select("*").execute()
    datos = respuesta.data
    
    if datos:
        df = pd.DataFrame(datos)
        
        # Crear TRES pestañas ahora
        tab1, tab2, tab3 = st.tabs(["👥 Gestión de Invitados", "📊 Presupuesto Profesional", "🤝 Financiamiento"])
        
        # ==========================================
        # PESTAÑA 1: GESTIÓN DE INVITADOS
        # ==========================================
        with tab1:
            df['nombre_completo'] = (df['nombre'].str.strip() + " " + df['apellido'].str.strip()).str.lower()
            conteo_sobrantes = df.duplicated(subset=['nombre_completo'], keep='first').sum()
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric(label="Total de Registros (Actual)", value=len(df))
            with col2:
                st.metric(label="Registros Repetidos (Sobrantes)", value=conteo_sobrantes, delta_color="inverse")

            st.divider()

            if conteo_sobrantes > 0:
                st.warning(f"⚠️ Se han detectado {conteo_sobrantes} registros repetidos. Revisa la tabla a continuación:")
                df_mostrar_repetidos = df[df.duplicated(subset=['nombre_completo'], keep=False)].sort_values(by='nombre_completo')
                columnas_revisar = ['nombre', 'apellido', 'gusto_comida', 'disco_preferido']
                if 'created_at' in df.columns: columnas_revisar.append('created_at')
                
                st.dataframe(df_mostrar_repetidos[[c for c in columnas_revisar if c in df_mostrar_repetidos.columns]], use_container_width=True)
                
                if st.button("🗑️ CONFIRMAR Y ELIMINAR REPETIDOS"):
                    with st.spinner("Limpiando base de datos..."):
                        exitos, errores = eliminar_duplicados(df)
                        if exitos > 0:
                            st.success(f"¡Limpieza completada! Se eliminaron {exitos} registros.")
                            st.rerun()
                        if errores > 0:
                            st.error(f"Hubo problemas eliminando {errores} registros.")
            else:
                st.success("✅ La base de datos está limpia. No hay invitados repetidos.")

            st.divider()
            st.subheader("Lista General de Invitados")
            columnas_visibles = ['nombre', 'apellido', 'gusto_comida', 'disco_preferido']
            st.dataframe(df[[c for c in columnas_visibles if c in df.columns]].sort_values(by='nombre'), use_container_width=True)
            if st.button("🔄 Refrescar Lista"): st.rerun()

        # ==========================================
        # PESTAÑA 2: CALCULADORA PROFESIONAL
        # ==========================================
        with tab2:
            st.header("Estructura de Costos")
            
            # --- MEMORIA DEL SISTEMA: Cargar la última configuración si existe ---
            if 'config_guardada' not in st.session_state:
                st.session_state.config_guardada = {}
                try:
                    res_hist = supabase.table("presupuestos_historicos").select("datos_json").order("fecha_registro", desc=True).limit(1).execute()
                    if res_hist.data and res_hist.data[0].get('datos_json'):
                        st.session_state.config_guardada = res_hist.data[0]['datos_json']
                except:
                    pass
            
            conf = st.session_state.config_guardada

            # Contar preferencias reales
            conteo_comida = df['gusto_comida'].value_counts()
            total_carne = conteo_comida.get('Carne', 0) + conteo_comida.get('Sin preferencia', 0)
            total_pollo = conteo_comida.get('Pollo', 0)
            total_vege = conteo_comida.get('Vegetariano', 0)
            total_invitados = len(df)
            
            st.info(f"**Invitados Confirmados:** {total_carne} Carne | {total_pollo} Pollo | {total_vege} Veggie | **Total: {total_invitados} personas**")

            with st.form("form_presupuesto"):
                
                st.subheader("📈 COSTOS VARIABLES (Dependen del # de invitados)")
                col_c1, col_c2 = st.columns(2)
                precio_carne = col_c1.number_input("Precio por Kilo Carne ($)", value=int(conf.get('precio_carne', 35000)))
                gr_carne = col_c2.number_input("Gramos carne por persona", value=int(conf.get('gr_carne', 350)))
                
                col_p1, col_p2 = st.columns(2)
                precio_pollo = col_p1.number_input("Precio por Kilo Pollo ($)", value=int(conf.get('precio_pollo', 18000)))
                gr_pollo = col_p2.number_input("Gramos pollo por persona", value=int(conf.get('gr_pollo', 350)))
                
                st.markdown("---")
                st.write("**Pasabocas, Entradas y Acompañamientos**")
                nombre_pasa = st.text_input("Tipo de Pasabocas", value=conf.get('nombre_pasa', "Deditos y empanaditas"))
                col_b1, col_b2, col_b3 = st.columns(3)
                precio_unid_pasa = col_b1.number_input("Precio Unidad Pasaboca ($)", value=int(conf.get('precio_unid_pasa', 1200)))
                unid_persona_ronda = col_b2.number_input("Und. por persona (por ronda)", value=int(conf.get('unid_persona_ronda', 4)))
                num_rondas = col_b3.number_input("Número de rondas", value=int(conf.get('num_rondas', 2)))
                
                col_a1, col_a2 = st.columns(2)
                costo_entrada_pp = col_a1.number_input("Costo Entrada p/p ($)", value=int(conf.get('costo_entrada_pp', 8000)))
                costo_acompa = col_a2.number_input("Acompañamientos p/p ($) (Yuca, papa, etc.)", value=int(conf.get('costo_acompa', 6000)))
                
                presu_bebidas = st.number_input("Presupuesto Bebidas e Hielo ($)", value=int(conf.get('presu_bebidas', 250000)))
                
                st.divider()

                st.subheader("📌 COSTOS FIJOS (Independientes de la asistencia)")
                col_f1, col_f2 = st.columns(2)
                costo_lugar = col_f1.number_input("Alquiler del Lugar / Rancho JP ($)", value=int(conf.get('costo_lugar', 500000)), step=50000)
                costo_torta = col_f2.number_input("Costo de la Torta ($)", value=int(conf.get('costo_torta', 150000)), step=10000)
                
                col_f3, col_f4 = st.columns(2)
                costo_dj = col_f3.number_input("DJ Calao y Sonido ($)", value=int(conf.get('costo_dj', 250000)), step=10000)
                costo_deco = col_f4.number_input("Decoración Neón y Ambientación ($)", value=int(conf.get('costo_deco', 100000)), step=10000)
                
                st.divider()
                margen = st.slider("Margen Financiero para Imprevistos (%)", 0, 20, int(conf.get('margen', 10)))
                
                generar = st.form_submit_button("💰 GENERAR INFORME FINANCIERO")

            # Si le da a generar, procesamos y guardamos todo en la memoria de la sesión
            if generar:
                k_carne = (total_carne * gr_carne) / 1000
                gasto_carne = k_carne * precio_carne
                k_pollo = (total_pollo * gr_pollo) / 1000
                gasto_pollo = k_pollo * precio_pollo
                
                total_unidades_pasa = total_invitados * unid_persona_ronda * num_rondas
                gasto_pasabocas = total_unidades_pasa * precio_unid_pasa
                gasto_entradas = total_invitados * costo_entrada_pp
                gasto_acompa = total_invitados * costo_acompa
                
                total_variables = gasto_carne + gasto_pollo + gasto_pasabocas + gasto_entradas + gasto_acompa + presu_bebidas
                total_fijos = costo_lugar + costo_torta + costo_dj + costo_deco
                subtotal = total_variables + total_fijos
                monto_imprevistos = subtotal * (margen / 100)
                gran_total = subtotal + monto_imprevistos
                cuota_por_persona = gran_total / total_invitados if total_invitados > 0 else 0
                
                # Guardamos los resultados y la configuración en el Session State
                st.session_state.resultados = {
                    "k_carne": k_carne, "gasto_carne": gasto_carne, "k_pollo": k_pollo, "gasto_pollo": gasto_pollo,
                    "total_unidades_pasa": total_unidades_pasa, "gasto_pasabocas": gasto_pasabocas,
                    "gasto_entradas": gasto_entradas, "gasto_acompa": gasto_acompa, 
                    "total_variables": total_variables, "total_fijos": total_fijos,
                    "subtotal": subtotal, "monto_imprevistos": monto_imprevistos, "gran_total": gran_total,
                    "cuota_por_persona": cuota_por_persona
                }
                
                st.session_state.config_guardada = {
                    "precio_carne": precio_carne, "gr_carne": gr_carne, "precio_pollo": precio_pollo, 
                    "gr_pollo": gr_pollo, "nombre_pasa": nombre_pasa, "precio_unid_pasa": precio_unid_pasa, 
                    "unid_persona_ronda": unid_persona_ronda, "num_rondas": num_rondas, 
                    "costo_entrada_pp": costo_entrada_pp, "costo_torta": costo_torta,
                    "costo_acompa": costo_acompa, "presu_bebidas": presu_bebidas, 
                    "costo_lugar": costo_lugar, "costo_dj": costo_dj, "costo_deco": costo_deco, "margen": margen
                }

            # ================= INTERFAZ DE RESULTADOS =================
            if 'resultados' in st.session_state:
                res = st.session_state.resultados
                conf = st.session_state.config_guardada
                
                st.divider()
                st.subheader("📊 Resumen Ejecutivo")
                
                r1, r2, r3, r4 = st.columns(4)
                r1.metric("Kilos Carne", f"{res['k_carne']:.1f} kg")
                r2.metric("Kilos Pollo", f"{res['k_pollo']:.1f} kg")
                r3.metric("Total Pasabocas", f"{res['total_unidades_pasa']} und")
                r4.metric("Cuota Est. p/p", f"${res['cuota_por_persona']:,.0f}")

                resumen_print = f"""
                ### 📝 REPORTE FINANCIERO: POOL PARTY SALVADOR
                ---
                **Asistencia Proyectada:** {total_invitados} personas
                
                **1. COSTOS VARIABLES:**
                * Proteína (Carne {res['k_carne']:.1f}kg + Pollo {res['k_pollo']:.1f}kg): **${(res['gasto_carne'] + res['gasto_pollo']):,.0f}**
                * Entradas: **${res.get('gasto_entradas', 0):,.0f}**
                * Pasabocas ({conf['nombre_pasa']} - {res['total_unidades_pasa']} unds): **${res['gasto_pasabocas']:,.0f}**
                * Acompañamientos: **${res['gasto_acompa']:,.0f}**
                * Bebidas e Hielo: **${conf['presu_bebidas']:,.0f}**
                * **Subtotal Variables: ${res['total_variables']:,.0f}**
                
                **2. COSTOS FIJOS:**
                * Alquiler de Locación (Rancho JP): **${conf['costo_lugar']:,.0f}**
                * Torta: **${conf.get('costo_torta', 0):,.0f}**
                * DJ Calao y Sistema de Sonido: **${conf['costo_dj']:,.0f}**
                * Decoración Neón y Ambientación: **${conf['costo_deco']:,.0f}**
                * **Subtotal Fijos: ${res['total_fijos']:,.0f}**
                
                **3. CONSOLIDADO:**
                * Subtotal Operativo: ${res['subtotal']:,.0f}
                * Fondo de Imprevistos ({conf['margen']}%): ${res['monto_imprevistos']:,.0f}
                
                ### **VALOR TOTAL DEL PROYECTO: ${res['gran_total']:,.0f}**
                """
                
                # ADICIÓN AUTOMÁTICA DE APORTANTES AL REPORTE SI EXISTEN
                if st.session_state.aportantes:
                    resumen_print += "\n---\n**4. FINANCIAMIENTO Y APORTES:**\n"
                    total_recogido = 0
                    for ap in st.session_state.aportantes:
                        resumen_print += f"* {ap['nombre']} ({ap['detalle']}): **${ap['monto']:,.0f}**\n"
                        total_recogido += ap['monto']
                    resumen_print += f"\n* **Total Recaudado: ${total_recogido:,.0f}**\n"
                    resumen_print += f"* **Faltante por cubrir: ${(res['gran_total'] - total_recogido):,.0f}**\n"
                
                resumen_print += f"\n---\n*Punto de equilibrio (Cuota sugerida por invitado): ${res['cuota_por_persona']:,.0f}*"
                
                st.markdown(resumen_print)
                
                st.button("🖨️ PREPARAR PARA IMPRIMIR (Ctrl + P)", on_click=lambda: st.info("Usa Ctrl + P (o Cmd + P en Mac) para guardar este reporte en PDF o imprimirlo."))

                st.divider()
                st.subheader("💾 Opciones de Exportación")
                col_d1, col_d2, col_d3 = st.columns(3)

                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=11)
                for linea in resumen_print.split('\n'):
                    texto_limpio = linea.replace('**', '').replace('### ', '').replace('---', '-'*50)
                    texto_limpio = texto_limpio.encode('latin-1', 'ignore').decode('latin-1').strip()
                    if texto_limpio:
                        pdf.cell(0, 7, txt=texto_limpio, ln=True)
                
                try:
                    pdf_bytes = pdf.output(dest='S').encode('latin-1', 'ignore')
                except AttributeError:
                    pdf_bytes = bytes(pdf.output())

                col_d1.download_button(label="📄 Descargar como PDF", data=pdf_bytes, file_name="Presupuesto_PoolParty.pdf", mime="application/pdf")

                doc = Document()
                doc.add_heading('Reporte Financiero Pool Party', 0)
                for linea in resumen_print.split('\n'):
                    texto_limpio = linea.replace('**', '').replace('### ', '').replace('---', '-'*50).replace('📝', '').strip()
                    if texto_limpio:
                        doc.add_paragraph(texto_limpio)
                b = BytesIO()
                doc.save(b)
                word_bytes = b.getvalue()

                col_d2.download_button(label="📝 Descargar como Word", data=word_bytes, file_name="Presupuesto_PoolParty.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                datos_guardar = {
                    "fecha_registro": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "total_invitados": total_invitados,
                    "costos_variables": float(res['total_variables']),
                    "costos_fijos": float(res['total_fijos']),
                    "gran_total": float(res['gran_total']),
                    "cuota_por_persona": float(res['cuota_por_persona']),
                    "datos_json": conf 
                }

                if col_d3.button("☁️ Guardar Presupuesto en Supabase"):
                    try:
                        supabase.table("presupuestos_historicos").insert(datos_guardar).execute()
                        st.success("✅ ¡Guardado! La próxima vez que entres, los precios se cargarán solos.")
                        st.balloons()
                    except Exception as e:
                        st.error(f"⚠️ Necesitas agregar una columna llamada 'datos_json' (tipo JSONB) en la tabla 'presupuestos_historicos'. Error: {e}")

        # ==========================================
        # PESTAÑA 3: FINANCIAMIENTO 
        # ==========================================
        with tab3:
            st.header("🤝 Gestión de Aportes y Financiamiento")
            
            if 'resultados' not in st.session_state:
                st.warning("⚠️ Primero debes ir a la pestaña 'Presupuesto Profesional' y darle al botón de Generar Informe para conocer el costo total de la fiesta.")
            else:
                res = st.session_state.resultados
                conf = st.session_state.config_guardada
                gran_total = res['gran_total']
                
                st.metric("Meta de Recaudo (Costo Total de la Fiesta)", f"${gran_total:,.0f}")
                st.divider()

                with st.form("form_aportes"):
                    st.subheader("Asignar un Patrocinador o Responsable")
                    
                    c_resp1, c_resp2 = st.columns(2)
                    nombre_aportante = c_resp1.text_input("Nombre de la Persona (Ej. Mamá, Yo, Tío Juan)")
                    
                    # AQUÍ ESTÁ LA MAGIA DEL SALDO FALTANTE
                    tipo_aporte = c_resp2.selectbox("Tipo de Aporte", [
                        "Asumir un Gasto Específico", 
                        "Monto Fijo ($)", 
                        "Porcentaje del Total General (%)",
                        "Porcentaje del Saldo Faltante (%)"
                    ])
                    
                    st.markdown("---")
                    st.write("**Detalles del Aporte**")
                    
                    c_det1, c_det2, c_det3 = st.columns(3)
                    
                    gastos_disponibles = {
                        "Proteína (Carne y Pollo)": res['gasto_carne'] + res['gasto_pollo'],
                        "Entradas": res.get('gasto_entradas', 0),
                        f"Pasabocas ({conf['nombre_pasa']})": res['gasto_pasabocas'],
                        "Acompañamientos": res['gasto_acompa'],
                        "Bebidas e Hielo": conf['presu_bebidas'],
                        "Alquiler de Locación": conf['costo_lugar'],
                        "Torta": conf.get('costo_torta', 0),
                        "DJ y Sonido": conf['costo_dj'],
                        "Decoración": conf['costo_deco'],
                        "Fondo de Imprevistos": res['monto_imprevistos']
                    }
                    
                    item_especifico = c_det1.selectbox("Si eligió 'Gasto Específico', ¿Qué va a pagar?", ["Seleccionar..."] + list(gastos_disponibles.keys()))
                    monto_fijo = c_det2.number_input("Si eligió 'Monto Fijo', escribe el valor ($)", value=0, step=50000)
                    porcentaje_fijo = c_det3.number_input("Si eligió 'Porcentaje', escribe el %", value=0, max_value=100)
                    
                    agregar_aporte = st.form_submit_button("➕ AGREGAR AL RECAUDO")
                    
                if agregar_aporte:
                    if not nombre_aportante:
                        st.error("Por favor escribe el nombre de la persona.")
                    else:
                        valor_calculado = 0
                        detalle_texto = ""
                        
                        # Cálculos basados en la selección
                        if tipo_aporte == "Asumir un Gasto Específico" and item_especifico != "Seleccionar...":
                            valor_calculado = gastos_disponibles[item_especifico]
                            detalle_texto = f"Pago de {item_especifico}"
                            
                        elif tipo_aporte == "Monto Fijo ($)" and monto_fijo > 0:
                            valor_calculado = monto_fijo
                            detalle_texto = "Aporte de monto fijo"
                            
                        elif tipo_aporte == "Porcentaje del Total General (%)" and porcentaje_fijo > 0:
                            valor_calculado = gran_total * (porcentaje_fijo / 100)
                            detalle_texto = f"Aporte del {porcentaje_fijo}% del total general"
                            
                        elif tipo_aporte == "Porcentaje del Saldo Faltante (%)" and porcentaje_fijo > 0:
                            # Calcular cuánto falta ANTES de agregar este nuevo aporte
                            recogido_hasta_ahora = sum(ap['monto'] for ap in st.session_state.aportantes)
                            saldo_actual = gran_total - recogido_hasta_ahora
                            
                            if saldo_actual <= 0:
                                st.error("¡El presupuesto ya está cubierto en su totalidad! No hay saldo faltante.")
                                valor_calculado = 0
                            else:
                                valor_calculado = saldo_actual * (porcentaje_fijo / 100)
                                detalle_texto = f"Aporte del {porcentaje_fijo}% del saldo faltante"
                                
                        else:
                            st.error("Verifica los datos del aporte. Selecciona un item, un monto o un porcentaje mayor a 0.")
                            valor_calculado = 0
                            
                        if valor_calculado > 0:
                            st.session_state.aportantes.append({
                                "nombre": nombre_aportante,
                                "detalle": detalle_texto,
                                "monto": valor_calculado
                            })
                            st.success(f"¡Aporte de {nombre_aportante} agregado exitosamente!")
                            st.rerun()

                # --- MOSTRAR EL BALANCE FINANCIERO ---
                if st.session_state.aportantes:
                    st.subheader("📈 Balance Financiero")
                    df_aportes = pd.DataFrame(st.session_state.aportantes)
                    total_recogido = df_aportes['monto'].sum()
                    faltante = gran_total - total_recogido
                    
                    c_bal1, c_bal2 = st.columns(2)
                    c_bal1.metric("Dinero Asegurado", f"${total_recogido:,.0f}")
                    
                    if faltante > 0:
                        c_bal2.metric("Faltante por Cubrir", f"${faltante:,.0f}", delta="Aún falta", delta_color="inverse")
                    elif faltante < 0:
                         c_bal2.metric("Sobrante a favor", f"${abs(faltante):,.0f}", delta="¡Superaron la meta!", delta_color="normal")
                    else:
                        c_bal2.metric("Faltante por Cubrir", "$0", delta="¡Presupuesto Exacto Cubierto!", delta_color="normal")
                    
                    # Barra de progreso visual
                    progreso = min(total_recogido / gran_total, 1.0)
                    st.progress(progreso)
                    
                    st.write("**Lista de Responsables:**")
                    
                    # Formato moneda para la tabla para que se vea igual que en la imagen que mandaste
                    df_aportes_mostrar = df_aportes.copy()
                    df_aportes_mostrar['monto'] = df_aportes_mostrar['monto'].apply(lambda x: f"${x:,.0f}")
                    st.dataframe(df_aportes_mostrar, use_container_width=True)
                    
                    if st.button("🗑️ Limpiar Lista de Aportes"):
                        st.session_state.aportantes = []
                        st.rerun()

    else:
        st.info("Aún no hay invitados registrados en la base de datos.")
        if st.button("🔄 Reintentar"): st.rerun()
        
except Exception as e:
    st.error(f"Error de conexión o de datos: {e}")